from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


ROOT_DIR = Path(r"D:\Web_QL_ChiTieu")
TEST_DIR = ROOT_DIR / "test"
OUTPUT_PATH = TEST_DIR / "BAO_CAO_KIEM_THU_FAIL_FIRST_CAP_NHAT.docx"

SCREENSHOTS = {
    "incident_01_fail": Path(
        r"C:\Users\LAN ANH\Pictures\Screenshots\Screenshot 2026-05-08 150618.png"
    ),
    "incident_02_fail": Path(
        r"C:\Users\LAN ANH\Pictures\Screenshots\Screenshot 2026-05-08 150816.png"
    ),
    "incident_03_fail": Path(
        r"C:\Users\LAN ANH\Pictures\Screenshots\Screenshot 2026-05-08 152555.png"
    ),
    "all_fail": Path(
        r"C:\Users\LAN ANH\Pictures\Screenshots\Screenshot 2026-05-08 152631.png"
    ),
    "incident_03_pass": Path(
        r"C:\Users\LAN ANH\Pictures\Screenshots\Screenshot 2026-05-08 152919.png"
    ),
    "incident_01_pass": Path(
        r"C:\Users\LAN ANH\Pictures\Screenshots\Screenshot 2026-05-08 152938.png"
    ),
    "incident_02_pass": Path(
        r"C:\Users\LAN ANH\Pictures\Screenshots\Screenshot 2026-05-08 153508.png"
    ),
    "all_pass": Path(
        r"C:\Users\LAN ANH\Pictures\Screenshots\Screenshot 2026-05-08 153543.png"
    ),
    "lint_pass": Path(
        r"C:\Users\LAN ANH\Pictures\Screenshots\Screenshot 2026-05-08 153628.png"
    ),
}


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(15)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    sizes = {1: 13, 2: 12, 3: 12}
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(sizes.get(level, 12))


def add_line(document: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_blank(document: Document) -> None:
    document.add_paragraph("")


def add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_picture_with_caption(
    document: Document, image_path: Path, caption: str, *, width: float = 6.2
) -> None:
    if not image_path.exists():
        add_line(document, f"[Thiếu ảnh: {image_path}]", italic=True)
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(11)


def add_cover(document: Document) -> None:
    add_blank(document)
    add_blank(document)
    add_blank(document)
    add_title(document, "BÁO CÁO KIỂM THỬ HỆ THỐNG THEO HƯỚNG FAIL FIRST")
    add_blank(document)
    add_line(document, "Dự án: Web quản lý chi tiêu", bold=True)
    add_line(document, "Phạm vi: Frontend React/Vite, Backend Node.js, API local, kiểm thử cấu hình và validation.")
    add_line(document, "Số incident sử dụng trong báo cáo: 03")
    add_line(document, "Nguyên tắc thực hiện: phải tái hiện lỗi và chạy ra FAIL trước, sau đó mới sửa để PASS.")
    add_line(document, "Kết quả cuối cùng sau khi sửa: 11/11 test PASS và lint frontend PASS.")
    add_page_break(document)


def add_overview(document: Document) -> None:
    add_heading(document, "1. Mục tiêu kiểm thử")
    add_line(
        document,
        "Báo cáo này tập trung vào 3 incident chính của hệ thống local full stack. Mỗi incident đều được xử lý theo cùng một quy trình: tái hiện lỗi, chạy test để ghi nhận FAIL, phân tích nguyên nhân, áp dụng hướng giải quyết, sau đó chạy lại để xác nhận PASS.",
    )

    add_heading(document, "2. Công cụ và lệnh sử dụng")
    add_bullet(document, "Chạy hệ thống local: npm run dev")
    add_bullet(document, "Chạy incident 01: npm run test:incident:01")
    add_bullet(document, "Chạy incident 02: npm run test:incident:02")
    add_bullet(document, "Chạy incident 03: npm run test:incident:03")
    add_bullet(document, "Chạy toàn bộ test: npm test")
    add_bullet(document, "Kiểm tra frontend lint: npm --prefix Frontend run lint")

    add_heading(document, "3. Quy trình Fail First")
    add_number(document, "Chủ động tạo trạng thái lỗi trong cấu hình hoặc logic.")
    add_number(document, "Chạy test để hệ thống phát sinh FAIL.")
    add_number(document, "Phân tích nguyên nhân gốc và xác định lớp xảy ra lỗi.")
    add_number(document, "Sửa cấu hình hoặc code.")
    add_number(document, "Chạy lại test để xác nhận PASS.")
    add_page_break(document)


def add_incident_01(document: Document) -> None:
    add_heading(document, "INCIDENT 01: Frontend local trỏ sai backend local")
    add_line(document, "Mô tả sự cố:", bold=True)
    add_line(
        document,
        "Biến VITE_API_URL trong Frontend/.env bị cấu hình trỏ sang backend deploy trên Vercel thay vì http://localhost:3000/api. Khi đó giao diện chạy local nhưng dữ liệu không đi qua backend local nên kết quả test bị sai môi trường.",
    )

    add_line(document, "Cách tái hiện:", bold=True)
    add_number(document, "Sửa Frontend/.env thành VITE_API_URL=https://dev-ops-mini-biun.vercel.app hoặc endpoint deploy tương đương.")
    add_number(document, "Giữ backend và frontend local đang chạy.")
    add_number(document, "Chạy lệnh npm run test:incident:01.")

    add_line(document, "Minh chứng FAIL:", bold=True)
    add_picture_with_caption(
        document,
        SCREENSHOTS["incident_01_fail"],
        "Hình 1. Incident 01 ở trạng thái FAIL: test phát hiện Frontend/.env không trỏ về http://localhost:3000/api.",
    )

    add_line(document, "Nguyên nhân:", bold=True)
    add_line(
        document,
        "Frontend đang dùng endpoint deploy thay vì endpoint local. Điều này làm mất tính nhất quán của môi trường kiểm thử và khiến việc debug theo code local không còn chính xác.",
    )

    add_line(document, "Hướng giải quyết:", bold=True)
    add_number(document, "Mở Frontend/.env.")
    add_number(document, "Đổi giá trị thành VITE_API_URL=http://localhost:3000/api.")
    add_number(document, "Chạy lại npm run test:incident:01 để xác nhận.")

    add_line(document, "Minh chứng PASS:", bold=True)
    add_picture_with_caption(
        document,
        SCREENSHOTS["incident_01_pass"],
        "Hình 2. Incident 01 sau khi sửa: test PASS, frontend local đã trỏ đúng backend local.",
    )
    add_page_break(document)


def add_incident_02(document: Document) -> None:
    add_heading(document, "INCIDENT 02: Backend CORS không cho phép frontend local")
    add_line(document, "Mô tả sự cố:", bold=True)
    add_line(
        document,
        "CORS_ORIGIN trong Backend/.env bị đặt thành domain deploy, khiến backend local trả Access-Control-Allow-Origin sai. Trình duyệt sẽ chặn frontend local khi gửi request API.",
    )

    add_line(document, "Cách tái hiện:", bold=True)
    add_number(document, "Sửa Backend/.env thành CORS_ORIGIN=https://dev-o1.vercel.app.")
    add_number(document, "Restart lại npm run dev để backend nạp env mới.")
    add_number(document, "Chạy lệnh npm run test:incident:02.")

    add_line(document, "Minh chứng FAIL:", bold=True)
    add_picture_with_caption(
        document,
        SCREENSHOTS["incident_02_fail"],
        "Hình 3. Incident 02 ở trạng thái FAIL: test phát hiện CORS backend không khớp với origin của frontend local.",
    )

    add_line(document, "Nguyên nhân:", bold=True)
    add_line(
        document,
        "Backend local vẫn trả header CORS theo domain deploy. Dù frontend local có lên giao diện, các request API thật vẫn bị trình duyệt từ chối do sai origin.",
    )

    add_line(document, "Hướng giải quyết:", bold=True)
    add_number(document, "Mở Backend/.env.")
    add_number(document, "Đổi CORS_ORIGIN thành http://localhost:5173.")
    add_number(document, "Dừng và chạy lại npm run dev để backend nạp lại cấu hình.")
    add_number(document, "Chạy lại npm run test:incident:02.")

    add_line(document, "Minh chứng PASS:", bold=True)
    add_picture_with_caption(
        document,
        SCREENSHOTS["incident_02_pass"],
        "Hình 4. Incident 02 sau khi sửa: cả 2 test CORS đều PASS.",
    )
    add_page_break(document)


def add_incident_03(document: Document) -> None:
    add_heading(document, "INCIDENT 03: Backend trả 500 thay vì 400 cho lỗi validation")
    add_line(document, "Mô tả sự cố:", bold=True)
    add_line(
        document,
        "Khi gửi dữ liệu giao dịch không hợp lệ như thiếu title, amount nhỏ hơn hoặc bằng 0, date sai định dạng hoặc body JSON lỗi, backend đang trả mã 500 thay vì 400. Đây là phân loại sai bản chất lỗi.",
    )

    add_line(document, "Cách tái hiện:", bold=True)
    add_number(document, "Tạm làm backend gom lỗi validation vào nhánh trả 500.")
    add_number(document, "Restart lại backend bằng npm run dev.")
    add_number(document, "Chạy lệnh npm run test:incident:03.")

    add_line(document, "Minh chứng FAIL:", bold=True)
    add_picture_with_caption(
        document,
        SCREENSHOTS["incident_03_fail"],
        "Hình 5. Incident 03 ở trạng thái FAIL: 4 test validation đều nhận 500 thay vì 400.",
    )

    add_line(document, "Nguyên nhân:", bold=True)
    add_line(
        document,
        "Lỗi đầu vào của người dùng bị xử lý như lỗi hệ thống nội bộ. Điều này làm frontend khó hiển thị thông báo đúng và gây khó khăn cho việc phân biệt lỗi nghiệp vụ với lỗi server.",
    )

    add_line(document, "Hướng giải quyết:", bold=True)
    add_number(document, "Tại Backend/src/server.js, giữ riêng lớp ApiError cho lỗi validation.")
    add_number(document, "Bảo đảm các case input sai trả statusCode=400.")
    add_number(document, "Chạy lại npm run test:incident:03 để xác nhận.")

    add_line(document, "Minh chứng PASS:", bold=True)
    add_picture_with_caption(
        document,
        SCREENSHOTS["incident_03_pass"],
        "Hình 6. Incident 03 sau khi sửa: toàn bộ 4 test validation đều PASS.",
    )
    add_page_break(document)


def add_summary(document: Document) -> None:
    add_heading(document, "TỔNG HỢP KẾT QUẢ KIỂM THỬ")
    add_line(document, "Trạng thái FAIL tổng khi 3 incident cùng được tái hiện:", bold=True)
    add_picture_with_caption(
        document,
        SCREENSHOTS["all_fail"],
        "Hình 7. Kết quả tổng khi chưa sửa: npm test ghi nhận 7 failed, 4 passed.",
    )

    add_line(document, "Trạng thái PASS tổng sau khi hoàn tất sửa lỗi:", bold=True)
    add_picture_with_caption(
        document,
        SCREENSHOTS["all_pass"],
        "Hình 8. Kết quả tổng sau khi sửa: npm test ghi nhận 11 passed.",
    )

    add_line(document, "Kiểm tra chất lượng frontend:", bold=True)
    add_picture_with_caption(
        document,
        SCREENSHOTS["lint_pass"],
        "Hình 9. Frontend lint PASS, không còn lỗi ESLint trong mã nguồn hiện tại.",
        width=5.5,
    )

    add_heading(document, "Kết luận")
    add_bullet(document, "Báo cáo đã giữ đúng 3 incident theo yêu cầu.")
    add_bullet(document, "Mỗi incident đều có đủ phần FAIL, phân tích nguyên nhân, hướng giải quyết và PASS.")
    add_bullet(document, "Kết quả cuối cùng của hệ thống local là 11/11 test PASS.")
    add_bullet(document, "Frontend lint PASS, phù hợp để tiếp tục kiểm thử nghiệp vụ hoặc chuẩn bị deploy.")


def build_document() -> Document:
    document = Document()
    set_default_font(document)
    add_cover(document)
    add_overview(document)
    add_incident_01(document)
    add_incident_02(document)
    add_incident_03(document)
    add_summary(document)
    return document


def save_document(document: Document) -> Path:
    try:
        document.save(OUTPUT_PATH)
        return OUTPUT_PATH
    except PermissionError:
        fallback_path = OUTPUT_PATH.with_name(
            f"{OUTPUT_PATH.stem}_MOI{OUTPUT_PATH.suffix}"
        )
        document.save(fallback_path)
        return fallback_path


if __name__ == "__main__":
    missing = [str(path) for path in SCREENSHOTS.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("Thiếu ảnh đầu vào:\n" + "\n".join(missing))

    document = build_document()
    saved_path = save_document(document)
    print(saved_path)
